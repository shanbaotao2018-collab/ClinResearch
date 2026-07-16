from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


fonts = ['Noto Sans CJK SC', 'Arial Unicode MS', 'SimSun', 'Hiragino Sans GB', 'PingFang SC']
doc = Document()
for font in fonts:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f'{font}: 中文字体测试 - 医学文献综述智能体')
    run.font.name = font
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:ascii'), font)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run._element.rPr.rFonts.set(qn('w:cs'), font)
    lang = OxmlElement('w:lang')
    run._element.rPr.append(lang)
    lang.set(qn('w:val'), 'zh-CN')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    lang.set(qn('w:bidi'), 'zh-CN')
    run._element.rPr.set(qn('w:rsidRPr'), '00000000')
doc.save('tmp/font_probe.docx')
