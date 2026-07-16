from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / '附件1：“智行合e，全员AI秀”AI应用技能大赛报名表.docx'
FORM_OUTPUT = ROOT / '附件1：“智行合e，全员AI秀”AI应用技能大赛报名表-医学科研智能体工作台预填版.docx'
BRIEF_OUTPUT = ROOT / '医学科研智能体工作台-报名项目说明.docx'

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'


def set_run_font(run, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_widths(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width * 1440)))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill: str = LIGHT_GRAY) -> None:
    table.style = 'Table Grid'
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, 10, INK, True)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, 9.5, RGBColor(0, 0, 0), False)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = paragraph.add_run(text)
    set_run_font(run, 15 if level == 1 else 12, RGBColor(46, 116, 181) if level == 1 else RGBColor(31, 77, 120), True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, 10.5, RGBColor(0, 0, 0), True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run, 10.5, RGBColor(0, 0, 0), False)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 10.5, RGBColor(0, 0, 0), False)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style='List Bullet')
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, RGBColor(0, 0, 0), False)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F9')
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(f'{label}  ')
    set_run_font(run, 10.5, INK, True)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, RGBColor(0, 0, 0), False)


def build_prefilled_form() -> None:
    doc = Document(TEMPLATE)
    table = doc.tables[0]
    table.cell(0, 1).text = '研启智研：医学科研智能体工作台'
    table.cell(1, 1).text = '技术研发与智能开发（推荐）'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 10.5, RGBColor(0, 0, 0), run.bold)
    doc.save(FORM_OUTPUT)


def build_project_brief() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run('“智行合e・全员AI秀”AI应用技能大赛 | 项目简介')
    set_run_font(header_run, 8.5, MUTED, False)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('医学科研智能体工作台 | 报名阶段项目说明')
    set_run_font(footer_run, 8.5, MUTED, False)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run('研启智研：医学科研智能体工作台')
    set_run_font(run, 22, INK, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run('基于 OpenCode + Research Skills 的可编排、可复核科研智能体')
    set_run_font(run, 11.5, MUTED, False)

    metadata = doc.add_table(rows=2, cols=2)
    set_table_widths(metadata, [1.55, 4.95])
    metadata.cell(0, 0).text = '参赛方向'
    metadata.cell(0, 1).text = '技术研发与智能开发'
    metadata.cell(1, 0).text = '交付计划'
    metadata.cell(1, 1).text = '报名阶段；计划于 7 月底交付可运行的科研工作流版本'
    for row in metadata.rows:
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, 10, RGBColor(0, 0, 0), cell == row.cells[0])

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_callout(doc, '一句话概述', '将分散的医学科研方法 Skills 组织为可执行的多智能体工作流，覆盖“问题设计—文献证据—结构化抽取—科研写作—质量复核”，并把关键过程沉淀为可追溯项目资产。')

    add_heading(doc, '一、业务痛点与参赛目标')
    add_bullet(doc, '医学科研任务跨越选题、检索、筛选、信息抽取、写作与质量核查，科研人员需要在多个工具和表格之间反复切换。')
    add_bullet(doc, '通用对话式 AI 能回答问题，却很难按科研规范固定流程、调用真实数据库、保存中间结果并支持人工复核。')
    add_bullet(doc, '目标是交付一套可在 OpenCode 中直接调用的医学科研智能体：让专业 Skills、外部工具和项目数据形成闭环。')

    add_heading(doc, '二、总体方案与技术路线')
    add_body(doc, '以 OpenCode 为智能体宿主和编排层，以 medical-research-skills 为专业方法库，以 MCP 为真实工具接入层，以 FastAPI/SQLite 为项目事实与审计底座；采用“编排型 Agent + 专业子 Agent + 固定规则流 + 人工确认”的混合模式。')
    flow = doc.add_table(rows=1, cols=5)
    set_table_widths(flow, [1.30, 1.30, 1.30, 1.30, 1.30])
    labels = ['问题设计\nPICO/方案', '证据检索\n双库与引文', '筛选抽取\n结构化证据', '综合写作\n综述/大纲', '质控复核\n留痕导出']
    for index, label in enumerate(labels):
        flow.cell(0, index).text = label
        set_cell_shading(flow.cell(0, index), LIGHT_BLUE)
        p = flow.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, 9.5, INK, True)

    add_heading(doc, '三、月底交付的能力范围', 1)
    capability = doc.add_table(rows=1, cols=3)
    set_table_widths(capability, [1.45, 3.0, 2.05])
    headers = ['能力模块', '月底交付内容', '形成的价值']
    for i, value in enumerate(headers):
        capability.cell(0, i).text = value
    rows = [
        ('科研问题与方案设计', '主 Agent 澄清研究问题，生成 PICO、研究类型建议、检索与执行计划。', '让模糊科研需求转为可执行任务。'),
        ('文献检索与筛选', '检索子 Agent 生成策略并调用 PubMed / Europe PMC；筛选子 Agent 输出初筛建议与理由。', '实现真实检索、去重、纳排建议和 PRISMA 计数。'),
        ('证据与方法学抽取', '调用临床研究信息与方法学抽取 Skills，形成研究设计、样本、干预和结局的结构化证据表。', '降低人工阅读和表格录入成本。'),
        ('科研写作辅助', '调用综述框架、方法写作与讨论架构 Skills，基于已核验项目数据生成结构化大纲和初稿。', '将证据整理衔接到规范化写作。'),
        ('质量与科研诚信复核', '调用引文追溯、撤稿监测与人工复核机制；对不确定信息标记 human_review。', '避免虚构引文和将 AI 建议误作最终结论。'),
        ('项目化交付与审计', '通过 MCP 保存项目、检索策略、题录、决策、证据表和导出工作包。', '使过程可追溯、可复用、可复核。'),
    ]
    for module, implementation, value in rows:
        cells = capability.add_row().cells
        cells[0].text = module
        cells[1].text = implementation
        cells[2].text = value
    style_table(capability)

    doc.add_page_break()
    add_heading(doc, '四、智能体与 Skills 封装方式')
    verification = doc.add_table(rows=1, cols=4)
    set_table_widths(verification, [1.55, 1.65, 1.65, 1.65])
    for i, value in enumerate(['层级', '构成', '职责', '状态']):
        verification.cell(0, i).text = value
    results = [
        ('主编排', 'research-orchestrator', '任务拆解、路由、结构化交付、确认点控制', '月底交付'),
        ('证据检索', 'search-agent', '检索式、数据库策略、真实检索与引文扩展', '已具备基础'),
        ('证据筛选', 'screening-agent', '题录初筛、纳排理由与人工复核分流', '已具备基础'),
        ('抽取与写作', 'extraction / writing roles', '结构化证据、综述框架、方法与讨论初稿', '月底补齐封装'),
        ('质量复核', 'quality rules + Skills', '撤稿核查、证据边界、审计和导出', '月底补齐封装'),
    ]
    for item in results:
        cells = verification.add_row().cells
        for index, value in enumerate(item):
            cells[index].text = value
    style_table(verification)
    add_body(doc, '现有基础：已完成 OpenCode 主/子 Agent 骨架、14 个精选 Research Skills 同步机制、9 个 MCP 工具，以及真实 PubMed / Europe PMC 检索与两类课题的检索筛选闭环验证。')

    add_heading(doc, '五、创新点与推广价值')
    add_bullet(doc, '不是单一聊天机器人，而是“Agent 编排 + Skills 方法论 + MCP 工具执行 + 项目数据留痕”的科研生产系统。')
    add_bullet(doc, 'Skills 按任务路由调用，而不是一次性堆叠：检索、筛选、抽取、写作、质控各有输入输出和责任边界。')
    add_bullet(doc, '把可重复的科研流程固化为模板，可在不同病种、研究问题和综述类型之间复用。')
    add_bullet(doc, '坚持人机协同与科研诚信：任何最终检索式、纳排标准、研究结论及投稿级内容均需人工确认。')

    add_heading(doc, '六、月底交付验收与能力边界')
    add_body(doc, '计划交付：可在 OpenCode 中运行的主编排 Agent；检索、筛选、抽取、写作/质控角色定义；精选 Skills 调用规则；真实检索 MCP 工具；统一结构化输出模板；2 至 3 个医学课题演示及可导出的项目工作包。')
    add_body(doc, '能力边界：本项目服务医学科研辅助，不提供诊断或治疗建议；AI 不替代科研人员做最终纳排、偏倚风险评价或临床结论。全文 PDF 解析、效应量自动抽取和投稿级系统综述将按验证结果逐步扩展，不在报名阶段承诺为已完成能力。')

    doc.save(BRIEF_OUTPUT)


if __name__ == '__main__':
    build_prefilled_form()
    build_project_brief()
    print(FORM_OUTPUT)
    print(BRIEF_OUTPUT)
