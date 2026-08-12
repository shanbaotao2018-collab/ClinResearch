"""Build a competition-ready local self-test record from verified project artifacts."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(
    "/Users/shanbaotao/Documents/agent 2/deliverables/"
    "佐证材料4-临研智策V0.1本地真实案例自测记录.docx"
)

NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F5F7"
DARK_GRAY = RGBColor(89, 89, 89)


def set_run_font(run, name="PingFang SC", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        set_run_font(paragraph.add_run(line), size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        if widths:
            cell.width = widths[index]
        shade(cell, NAVY.__str__().replace("'", "").replace("RGBColor(", "").replace(")", ""))
        set_cell_text(cell, header, bold=True, color=RGBColor(255, 255, 255), size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if widths:
                cells[index].width = widths[index]
            shade(cells[index], LIGHT_GRAY if len(table.rows) % 2 else "FFFFFF")
            set_cell_text(cells[index], value, size=8.8)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=15 if level == 1 else 11.5, bold=True, color=NAVY if level == 1 else BLUE)
    return paragraph


def add_body(document, text, emphasis_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    if emphasis_prefix and text.startswith(emphasis_prefix):
        set_run_font(paragraph.add_run(emphasis_prefix), bold=True, color=NAVY)
        set_run_font(paragraph.add_run(text[len(emphasis_prefix):]))
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text), size=10)
    return paragraph


def add_case(document, number, title, goal, input_data, execution, outcome, evidence, boundary):
    add_heading(document, f"{number}. {title}", level=2)
    add_table(
        document,
        ["项目", "记录"],
        [
            ("验证目标", goal),
            ("输入与数据边界", input_data),
            ("实际执行", execution),
            ("结果", outcome),
            ("本地留痕/证据", evidence),
            ("边界说明", boundary),
        ],
        [Cm(3.0), Cm(13.5)],
    )


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("临研智策 V0.1 本地真实案例自测记录  |  第 ")
    set_run_font(run, size=8.5, color=DARK_GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run(" 页")
    set_run_font(run, size=8.5, color=DARK_GRAY)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)

    # Cover
    document.add_paragraph().paragraph_format.space_after = Pt(40)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("临研智策 V0.1"), size=28, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("本地真实案例自测记录"), size=20, bold=True, color=BLUE)
    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_after = Pt(18)
    set_run_font(lead.add_run("“智行合e・全员AI秀”AI应用技能大赛佐证材料"), size=11, color=DARK_GRAY)
    add_table(
        document,
        ["项目", "说明"],
        [
            ("产品", "临研智策：多智能体临床科研工作台"),
            ("版本", "V0.1"),
            ("验证日期", "2026 年 7 月 22 日"),
            ("验证范围", "研究设计、文献检索综述、证据抽取系统评价、科研写作、B/S 工作台"),
            ("数据口径", "公开发表文献、公开开放全文、脱敏或聚合的演示假设；不含患者身份信息"),
        ],
        [Cm(3.6), Cm(12.9)],
    )
    add_body(document, "声明：本记录用于验证医疗科研辅助工具链的执行、留痕与安全门禁，不构成诊断、治疗建议、伦理批准、统计审核、系统评价定稿或正式科研结论。")

    document.add_page_break()
    add_heading(document, "一、验证目的与判定口径")
    add_body(document, "本次自测验证“Agent 编排 + Skill 方法规范 + MCP 真实执行 + 项目事实记录 + 人工复核/受控导出”的闭环是否可用。测试采用本地 OpenCode、FastAPI/MCP 服务、SQLite 项目记录和 React B/S 工作台。")
    add_table(
        document,
        ["判定项", "通过口径"],
        [
            ("真实执行", "Agent 或后端工具实际调用检索、计算、入库、导出等能力；不以模型文字说明替代。"),
            ("可追溯", "项目、工作流、题录、证据、草稿、审批和导出可在本地数据库/工作台中查询。"),
            ("方法门禁", "关键持久化与导出操作要求已签名的 Skill 执行回执；缺失时必须停止。"),
            ("人机协同", "研究设计、系统评价、科研写作等高风险导出需由操作者确认，Agent 不得自行批准。"),
            ("安全边界", "不输入患者身份信息；不展示 RCT 实际分配序列；不把工具输出表述为临床结论。"),
        ],
        [Cm(3.2), Cm(13.3)],
    )
    add_heading(document, "二、当次回归验证结果")
    add_table(
        document,
        ["检查项", "执行结果", "判定"],
        [
            ("后端自动化测试", "2026-07-22 执行 pytest：39 passed；3 条第三方框架弃用警告。", "通过"),
            ("Agent-Skill 合同检查", "执行 node scripts/opencode/test-agent-skill-contract.mjs，输出 contracts are enforced。", "通过"),
            ("B/S 工作台构建", "执行 npm run build：TypeScript 编译与 Vite 生产构建成功。", "通过"),
        ],
        [Cm(3.2), Cm(10.4), Cm(2.9)],
    )

    add_heading(document, "三、真实案例自测记录")
    add_body(document, "以下案例均来自本地已保存的运行记录或真实公开数据源验证。为避免把历史演示与当次回归混淆，案例中分别标注运行日期、项目/工作流标识及已知限制。")

    add_case(
        document,
        "案例 1",
        "研究设计 Agent：SGLT2 与心衰住院风险务实 RCT",
        "验证从研究问题到研究设计包、基础样本量、随机化计划、审批门禁和受控导出的完整链路。",
        "脱敏/聚合的演示性研究假设；不含患者身份信息。研究类型为 efficacy，报告规范为 SPIRIT/CONSORT。",
        "2026-07-16，OpenCode study-design-agent 实际调用 Skills 与本地 MCP；项目 ID 15，workflow run ID 602700a20b6644b9b535cc8809141173。",
        "项目状态 exported；两组率比较基础样本量总计 712；外部演示审批 approved。审批后才生成受保护随机表，Agent 返回 allocation_visible_to_agent=false，导出包不含实际分配序列。",
        "签名 Skill 回执 8 个；项目、审批和工作流事件保存在本地数据库。来源：docs/three-real-agent-runs-2026-07-16.md。",
        "样本量仅为演示性基础计算，正式方案、统计参数和随机化必须由研究者/统计师复核。",
    )

    add_case(
        document,
        "案例 2",
        "研究设计 Agent：AI 辅助 CT 肺结节诊断准确性研究",
        "验证非随机诊断研究的报告规范选择、方案结构化和受控导出，不错误生成 RCT 分配表。",
        "脱敏/聚合的演示性研究假设；不含患者身份信息。研究类型为 diagnostic，报告规范为 STARD。",
        "2026-07-16，OpenCode study-design-agent；项目 ID 16，workflow run ID 309696a67b304e42844015db8a8cae58。",
        "项目状态 exported；已完成研究设计、纳排标准、结局、基础样本量和审批。随机化不适用，导出工作包未产生随机化内容。",
        "签名 Skill 回执 7 个；项目与导出状态持久化。来源：docs/three-real-agent-runs-2026-07-16.md。",
        "诊断准确性研究的正式样本量和参考标准设计需采用相应方法学并由专家复核。",
    )

    add_case(
        document,
        "案例 3",
        "Skill 门禁负向验证：心衰再入院预后队列研究",
        "验证系统在必需方法 Skill 未真实执行时会停止，而不是让模型以提示词伪造完成。",
        "脱敏/聚合的演示性研究假设；研究类型 prognosis，目标规范 TRIPOD。",
        "2026-07-16，项目 ID 20，workflow run ID 95586d56088842c5aa1b3ed4556c3d20；流程在样本量步骤前校验回执。",
        "系统拒绝样本量计算并停止在 content_drafted。错误信息指出 sample-size-basic 缺少已验证的 OpenCode Skill 回执；未进入审批或导出。",
        "已保存 6 个签名 Skill 回执和阻断事件。来源：docs/three-real-agent-runs-2026-07-16.md。",
        "该案例的“未完成”是预期安全结果，证明当前系统不会把缺少方法执行依据的任务伪装成完成。",
    )

    add_case(
        document,
        "案例 4",
        "文献检索综述 Agent：AI 眼底影像筛查糖尿病视网膜病变",
        "验证真实医学数据库检索、项目创建、题录导入、去重、标题摘要筛选和 PRISMA 记录。",
        "公开发表文献；PICO 为糖尿病人群中 AI 眼底影像筛查与人工眼科医生判读/参考标准的诊断准确性比较。",
        "项目 ID 2；主 Agent 完成项目创建、检索策略与 search-agent 调用，实际调用 PubMed 与 Europe PMC；导入 3 篇真实候选研究并调用 screening-agent。",
        "去重移除 0 篇；3 篇暂定 include；PRISMA 为 identified 3、deduplicated 3、screened 3、included 3、excluded 0；项目状态 exported。",
        "本地 citation_id、筛选原因、PRISMA 与导出包均已保存。来源：docs/case-ai-dr-diagnostic-accuracy-mvp-output.md。",
        "外部模型长链路在提交筛选决定前两次超时，最后筛选提交和导出由同一后端 HTTP 接口完成。因此，业务引擎闭环已验证，但该次外部模型连接稳定性不作为完全验收结论。",
    )

    add_case(
        document,
        "案例 5",
        "证据抽取与系统评价 Agent：公开全文二分类 Meta",
        "验证公开全文入库、结构化字段抽取、RoB 2 初评、二分类随机效应 Meta、森林图与范围审批。",
        "公开 PMC 全文：RECOVERY 和 WHO Solidarity 两篇住院 COVID-19 成人患者研究；仅用于历史公开文献方法学测试。",
        "项目 ID 4，workflow run ID ee543310823d4642bc37d65fbdfaaeb9；通过 fulltext-fetcher 获取公开全文，保存基线/结局与 RoB 2 结构化记录，运行 28 天全因死亡的 RR 随机效应 Meta。",
        "成功生成 2 条全文证据行、1 次 Meta 与 SVG 森林图；合并工具结果 RR 1.09，95% CI 0.99–1.20，I² 0.0%，k=2；审批后成功渲染最终证据包。",
        "每个偏倚风险域保留 Methods/Results 定位；审批范围状态 scope_current=true。来源：docs/mvp-evidence-extraction-acceptance-2026-07-20.md。",
        "上述 Meta 仅是工具链验证输出，不构成用药或临床结论；抽取、偏倚风险、结局定义和模型适用性仍需研究者复核。",
    )

    add_case(
        document,
        "案例 6",
        "科研写作 Agent：研究方案与标书草稿",
        "验证 Agent 只读取已保存来源、记录 source_manifest、生成版本化草稿并经人工审批后导出。",
        "来源为已保存研究设计项目：study_design #15（SGLT2 方案）与 #16（AI 辅助 CT 诊断准确性方案）；不输入患者身份数据。",
        "2026-07-16，research-writing-agent 分别生成 protocol 草稿 #4 和 proposal 草稿 #5；工具链读取来源、启动写作、保存草稿、请求审批、查询状态并导出。",
        "两份草稿均 exported；分别保留 methods、discussion、outline 或 proposal 内容、source_manifest、limitations 与 unresolved_items。",
        "数据库记录 HMAC 签名 Skill 回执、OpenCode 会话与执行时间；proposal 分支额外校验 research-proposal-generator。来源：docs/two-research-writing-agent-runs-2026-07-16.md。",
        "导出物是待复核草稿，不是正式标书、投稿稿件、伦理批准或临床结论；来源或内容变化后需生成新版本并重新审批。",
    )

    document.add_page_break()
    add_heading(document, "四、执行链路与可复现入口")
    add_body(document, "四类 Agent 的手工测试输入和原生确认操作已整理在本地。运行前须在后端配置本地演示审批密钥，并在仓库根目录启动 OpenCode；审批密钥不进入模型提示词、截图、普通导出或本记录。")
    add_table(
        document,
        ["对象", "可复现输入/说明"],
        [
            ("研究设计", "testdata/mvp-four-agent/05-study-design-internal-approval-input.json；详见 docs/manual-test-plan-native-approval.md 第 1 节。"),
            ("文献检索综述", "testdata/mvp-four-agent/02-literature-review-input.json；详见同文档第 2 节。"),
            ("证据抽取", "使用已完成筛选的 review 项目；详见同文档第 3 节及 docs/mvp-evidence-extraction-acceptance-2026-07-20.md。"),
            ("科研写作", "基于已保存 study-design 或 evidence 来源；详见同文档第 4 节及 docs/two-research-writing-agent-runs-2026-07-16.md。"),
            ("工作台", "apps/research-workbench-frontend/；当前是只读 P0，用于查看项目、审计、审批状态和受控导出。"),
        ],
        [Cm(3.2), Cm(13.3)],
    )
    add_heading(document, "五、结论与当前边界")
    add_bullet(document, "本地 V0.1 已验证四个业务 Agent 的核心闭环，并保留真实执行、项目记录、Skill 回执和审批/导出审计。")
    add_bullet(document, "当次代码回归为后端 39 项测试通过、Agent-Skill 合同检查通过、B/S 工作台生产构建通过。")
    add_bullet(document, "系统明确限制为科研辅助：研究设计、筛选、偏倚风险、Meta 与写作内容均需研究者、统计师或相应专家复核。")
    add_bullet(document, "当前 P0 不宣称已完成院内 EMR/FHIR 对接、中文病历 NLP、统一身份认证、网页发起任务、网页审批、复杂生存分析、ROC、LASSO 或 PSM。")

    add_heading(document, "六、真实测试结果包索引")
    add_body(document, "与本记录配套的“佐证材料4-本地真实测试结果包-2026-07-22.zip”收录了当次重新执行的测试日志、可复现测试输入、历史真实运行记录及脱敏后的工具链/Skill 回执摘要。结果包不包含密钥、原始签名、OpenCode 会话标识、RCT 实际随机分配序列、患者身份信息或非公开数据。")
    add_table(
        document,
        ["结果包内容", "对应本记录内容"],
        [
            ("01-03 当次测试日志", "第二节：后端 39 项测试、Agent-Skill 合同检查、B/S 工作台生产构建。"),
            ("04-07 历史真实运行记录", "第三节：案例 1 至案例 6 的项目、工作流、结果与已知限制。"),
            ("08 Skill 回执脱敏汇总", "证明各案例的已执行 Skill 名称和回执数量，不暴露签名或会话标识。"),
            ("09 工具链实际结果脱敏摘要", "案例 5 的两篇公开全文、事件计数、Meta 工具输出与审批状态摘要。"),
            ("10 可复现测试输入", "第四节：四 Agent MVP 的脱敏/公开测试输入。"),
        ],
        [Cm(5.1), Cm(11.4)],
    )

    add_heading(document, "七、来源索引")
    sources = [
        "docs/three-real-agent-runs-2026-07-16.md",
        "docs/case-ai-dr-diagnostic-accuracy-mvp-output.md",
        "docs/mvp-evidence-extraction-acceptance-2026-07-20.md",
        "docs/two-research-writing-agent-runs-2026-07-16.md",
        "docs/manual-test-plan-native-approval.md",
        "testdata/mvp-four-agent/README.md",
    ]
    for source in sources:
        add_bullet(document, source)

    document.core_properties.title = "临研智策 V0.1 本地真实案例自测记录"
    document.core_properties.subject = "AI 应用技能大赛佐证材料"
    document.core_properties.author = "临研智策项目组"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
