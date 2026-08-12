"""Rewrite the final competition application according to the preliminary guide."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE_DIR = Path("/Users/shanbaotao/Documents/agent 2")
FINAL_DIR = BASE_DIR / "deliverables/参赛材料终版"
SOURCE = FINAL_DIR / "附件2-智行合e全员AI秀AI应用技能大赛申报表-临床科研智能体工作台-V0.1.docx"
OUTPUT = FINAL_DIR / "附件2-智行合e全员AI秀AI应用技能大赛申报表-临床科研智能体工作台-V0.1-按初赛指南修订版.docx"
REASON_DOCX = FINAL_DIR / "附件2申报表修正原因说明-按初赛指南.docx"
REASON_MD = FINAL_DIR / "附件2申报表修正原因说明-按初赛指南.md"


FONT = "STSong"


FIELDS = {
    "案例名称": "临床科研智能体工作台：基于多智能体的科研设计、综述、证据评价与写作助手",
    "个人/团队信息": "队长：施静\n成员：单保涛、刘胜军、王有闻",
    "参赛方向": "技术研发与智能开发",
    "应用场景简述": (
        "本应用面向临床科研人员在课题立项、文献综述、系统评价和科研写作中的高频工作。"
        "传统流程需要在检索数据库、Excel、统计工具、Word 文档和多人沟通之间反复切换，"
        "容易出现检索式不可复现、筛选依据分散、证据抽取缺少来源、AI 生成内容无法证明是否调用专业方法等问题。\n"
        "本项目将上述流程沉淀为四个可闭环智能体：研究设计 Agent、文献检索综述 Agent、证据抽取与系统评价 Agent、科研写作 Agent。"
        "科研人员输入研究问题或测试数据后，系统按固定科研流程调用工具、保存项目记录、生成可审阅成果，并在高风险节点要求人工确认。"
    ),
    "使用AI工具/智能体类型": (
        "采用“OpenCode 多智能体 + medical-research-skills + MCP 工具服务 + B/S 工作台”的组合式智能体方案，"
        "不是简单问答式大模型应用。\n"
        "1. 智能体编排：OpenCode 作为运行入口，定义 4 个主 Agent 和 2 个文献专业子 Agent，负责任务路由、流程控制、结构化输出和人工确认。\n"
        "2. 医疗科研 Skills：复用并筛选 30 项 medical-research-skills，覆盖 PICO、纳排标准、样本量、随机化、PubMed 检索、文献筛选、全文抽取、RoB 2/NOS/QUADAS-2、Meta 分析、方法学写作等能力。\n"
        "3. MCP/后端工具：FastAPI + MCP 提供真实 PubMed/Europe PMC 检索、题录导入去重、PRISMA 计数、样本量计算、证据表保存、二分类 Meta、审批、审计和导出。\n"
        "4. 产品界面：React B/S 工作台集中查看四类 Agent 的项目详情、运行状态、审批状态、测试结果和导出材料。"
    ),
    "改造前现状/痛点": (
        "改造前，临床科研辅助工作主要依赖人工经验和分散工具：\n"
        "1. 选题和研究设计不规范：PICO、纳排标准、结局指标、样本量假设、随机化方案往往散落在文档中，缺少统一结构和复核记录。\n"
        "2. 文献综述重复劳动重：研究者需要手工构造检索式、跨数据库检索、导入题录、去重、筛选并维护 PRISMA，过程耗时且难复现。\n"
        "3. 证据抽取和系统评价门槛高：基线、结局、偏倚风险、Meta 分析需要方法学知识，普通 AI 问答容易出现来源不清或结论过度表达。\n"
        "4. 写作草稿缺少来源约束：AI 生成文本如果不绑定项目事实和来源清单，评审时难判断哪些内容来自证据、哪些只是模型推断。\n"
        "5. 安全风险突出：科研场景涉及患者数据和临床结论，必须避免把敏感数据放入模型或把 AI 建议当成最终结论。"
    ),
    "落地方案": (
        "按照初赛“做出来且能用”的要求，本项目已完成 V0.1 本地 MVP，并形成可演示、可测试、可追溯的闭环。\n"
        "整体技术路线：Agent 负责任务编排，Skills 提供医学科研方法约束，MCP 工具负责真实执行和落库，B/S 工作台负责结果查看与交付展示。\n"
        "四个 Agent 的落地内容如下：\n"
        "1. 研究设计 Agent：创建 study-design 项目，生成研究问题、PICO、纳排标准、结局、基础样本量和随机化计划；通过 OpenCode 原生 Allow/Deny 完成人工确认，审批后才允许生成受保护随机化结果。\n"
        "2. 文献检索综述 Agent：调用 search-agent 生成检索策略，真实访问 PubMed/Europe PMC，导入题录、去重、生成标题摘要筛选建议、保存 PRISMA 计数并导出项目包。\n"
        "3. 证据抽取与系统评价 Agent：基于已纳入文献和公开全文，保存基线/结局字段、RoB 2 初评、二分类 Meta 结果和森林图，并通过审批门禁导出证据包。\n"
        "4. 科研写作 Agent：只读取已保存的研究设计或证据项目，生成 protocol、proposal、methods、discussion 等草稿，保留 source_manifest、limitations 和 unresolved_items。\n"
        "已配套完成：后端自动化测试、Agent-Skill 合同检查、真实案例自测记录、测试结果包、产品截图、演示视频、产品架构文档和 GitHub 源码仓库。"
    ),
    "落地成效或预测成效": (
        "已落地成效：当前已完成本地 V0.1 MVP，不是仅停留在 PPT 方案。2026 年 7 月已完成 6 类真实自测案例和一轮回归验证：后端自动化测试 39 项通过，Agent-Skill 合同检查通过，B/S 工作台生产构建通过。\n"
        "已验证能力包括：真实 PubMed/Europe PMC 检索、题录导入和去重、PRISMA 计数、公开全文入库、结构化证据抽取、RoB 2 初评、二分类 Meta 分析、研究设计样本量计算、受控随机化、科研写作草稿、审批与导出审计。\n"
        "量化佐证：本地真实测试结果包包含 3 份当次测试日志、6 类案例记录、脱敏 Skill 回执汇总、工具链结果摘要和 5 份可复现测试输入；源代码已提交 GitHub，形成可复查材料链。\n"
        "预测推广成效：在真实科室试点后，预计可将单个课题从“选题构思到研究方案初稿”的整理时间由数小时至 1 天压缩到 30-60 分钟；文献综述的检索式生成、题录去重和初筛记录可由手工跨工具操作转为一键项目化记录；证据抽取和写作阶段可显著减少重复录入、来源遗漏和版本混乱。上述推广成效需在后续科室试点中按任务样本继续量化。"
    ),
    "落地使用周期": (
        "2026 年 7 月启动并完成 V0.1 本地 MVP 开发、四智能体联调、真实案例验证和初赛佐证材料整理。"
        "当前使用周期为本地验证与参赛演示阶段，已连续用于本项目申报材料准备、演示视频录制、自测记录生成和测试结果打包。"
        "尚未接入医院生产环境或真实患者数据，后续可选择 1-2 个科室进行脱敏科研场景试点。"
    ),
    "可行性评估": (
        "技术可行：系统已完成可运行 MVP，后端测试、Agent 合同检查、前端构建均已通过；核心能力基于 OpenCode、MCP、FastAPI、React、PubMed/Europe PMC 和开源 medical-research-skills，可复现、可维护。\n"
        "推广可行：四个 Agent 均是小切口闭环能力，可按科室病种替换输入，不依赖特定专科；同一套流程可用于心衰、糖尿病、肺结节、感染性疾病等科研课题。\n"
        "资源需求：本地 MVP 可单机运行；正式推广需接入统一身份认证、权限审计、对象存储、PostgreSQL、日志归档和院内合规审批流程。若后续接入 EMR 或病历 NLP，必须另行完成数据授权、脱敏和安全评估。\n"
        "实施路径：初赛阶段展示本地 MVP；复赛/试点阶段选择脱敏公开课题验证医生实际使用效果；生产阶段再接入院内安全体系和科研数据治理平台。"
    ),
    "数据安全与合规说明": (
        "本项目严格按初赛安全合规要求处理数据：不使用公司内部生产数据库，不调用生产环境接口，不使用客户敏感数据、员工隐私数据或未脱敏患者数据进行训练或演示。\n"
        "当前测试数据来源包括：互联网公开医学文献、PubMed/Europe PMC 公开题录、PMC 公开全文、人工构造的脱敏/聚合研究假设和公开可复现测试输入。\n"
        "安全措施包括：患者身份信息禁止进入模型提示词、MCP 调用和普通日志；模型 API Key、审批 Key、Skill 回执签名 Key 不写入仓库、截图、视频或导出材料；RCT 实际随机分配序列存放在受保护位置，不展示给 Agent、B/S 工作台或普通导出包。\n"
        "合规边界：系统定位为临床科研辅助工具，不提供诊断、治疗建议或最终临床决策；样本量、筛选决定、偏倚风险、Meta 结果和写作草稿均标注为待研究者/统计师复核。"
    ),
    "附件清单": (
        "佐证材料1-四智能体演示视频：展示研究设计、文献检索综述、证据抽取系统评价、科研写作四类 Agent 操作流程。\n"
        "佐证材料2-产品截图：展示 OpenCode Agent 执行、B/S 工作台、项目详情和审批状态。\n"
        "佐证材料3-产品架构与业务闭环.pdf：说明总体架构、四 Agent 架构、MCP 工具和安全门禁。\n"
        "佐证材料4-临床科研智能体工作台V0.1本地真实案例自测记录.docx：记录 6 类真实案例和测试结论。\n"
        "佐证材料4-本地真实测试结果包-2026-07-22.zip：包含测试日志、真实运行记录、脱敏 Skill 回执摘要和可复现输入。\n"
        "临床科研智能体工作台-汇报.pptx：初赛展示材料。\n"
        "源码仓库：https://github.com/shanbaotao2018-collab/ClinResearch"
    ),
}


REASONS = [
    (
        "参赛方向从全量枚举改为单选“技术研发与智能开发”。",
        "指南要求不同方向突出不同评分重点；原表保留全部方向会削弱定位。项目实际包含 Agent 编排、MCP 工具、后端、前端和测试，更符合技术研发与智能开发。",
    ),
    (
        "案例名称强化“多智能体”和四类科研闭环。",
        "指南要求案例名称简洁并体现创新性。原名称偏泛，修订后能直接让评委看到产品形态和技术亮点。",
    ),
    (
        "应用场景从概括描述扩展为具体科研流程。",
        "产品品质维度关注真实场景和工作流程。修订版明确课题立项、检索筛选、系统评价、科研写作四类任务，便于评委判断不是临时拼装。",
    ),
    (
        "AI 工具类型突出“不是简单调用大模型”。",
        "创新性维度容易扣分点是仅调用 ChatGPT。修订版明确 OpenCode 多智能体、30 项医疗科研 Skills、MCP 工具、B/S 工作台和真实数据库调用。",
    ),
    (
        "痛点部分补充了可评审的具体问题。",
        "应用价值维度要求痛点真实具体。修订版把痛点拆为流程不规范、文献重复劳动、证据抽取门槛、写作来源不清、安全风险五类。",
    ),
    (
        "落地方案按四个 Agent 分别说明已实现功能。",
        "产品品质和创新性都要求证明功能完整、流程清晰。修订版逐个说明研究设计、文献综述、证据评价、科研写作的输入、执行和输出。",
    ),
    (
        "成效部分区分“已落地验证”和“预测推广成效”。",
        "指南要求注明已落地还是预测成效，并尽量量化。修订版写入 39 项测试、6 类真实案例、3 份测试日志、5 份可复现输入等已验证数据，同时将医生端效率提升表述为后续试点预测，避免过度承诺。",
    ),
    (
        "落地周期如实说明为本地 MVP 验证阶段。",
        "指南关注持续使用，但项目当前尚未进生产环境。修订版说明已用于申报材料准备、演示视频、自测和结果打包，并明确后续科室试点计划。",
    ),
    (
        "可行性补充推广路径和生产化资源需求。",
        "应用价值维度关注可推广。修订版说明可跨专科复用，并列明统一认证、审计、PostgreSQL、对象存储和数据治理等生产化条件。",
    ),
    (
        "安全合规改为一票否决口径。",
        "指南明确安全合规打 0 分直接淘汰。修订版明确不使用生产库、客户敏感数据、员工隐私和未脱敏患者数据，并说明密钥、随机分配序列、临床结论边界。",
    ),
    (
        "附件清单补齐真实测试结果包和源码仓库。",
        "产品品质维度要求演示视频、截图、源代码、真实运行记录。修订版把已有视频、截图、架构文档、自测记录、结果包、PPT 和 GitHub 全部索引到附件清单。",
    ),
]


def set_run_font(run, size=9.5, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_text(cell, text):
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.split("\n")):
        if i:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def rewrite_application():
    doc = Document(SOURCE)
    table = doc.tables[0]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label in FIELDS:
            set_cell_text(row.cells[1], FIELDS[label])
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=run.font.size.pt if run.font.size else 10.5, bold=bool(run.bold))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=run.font.size.pt if run.font.size else 9.5, bold=bool(run.bold))
    doc.core_properties.title = "临床科研智能体工作台 V0.1 初赛申报表修订版"
    doc.core_properties.subject = "按初赛申报材料准备指南修订"
    doc.core_properties.author = "临床科研智能体工作台项目组"
    doc.save(OUTPUT)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 11.5, bold=True, color=RGBColor(31, 78, 121))


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), size=10.5)


def build_reason_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    add_heading(doc, "附件2申报表修正原因说明", 1)
    add_para(doc, "依据《副本初赛申报材料准备指南》，初赛评分重点为产品品质 25 分、创新性 30 分、应用价值 35 分、安全合规 10 分。其中应用价值与创新性合计 65 分，是材料需要重点强化的部分。本次修订遵循项目 V0.1 的真实完成情况，不把尚未生产化的能力表述为已上线。")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["修正项", "修正原因", "对应评分维度"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    score_map = [
        "整体印象/创新性",
        "整体印象/创新性",
        "产品品质",
        "创新性",
        "应用价值",
        "产品品质+创新性",
        "应用价值",
        "应用价值",
        "应用价值",
        "安全合规",
        "产品品质",
    ]
    for (item, reason), score in zip(REASONS, score_map):
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], reason)
        set_cell_text(cells[2], score)
    add_heading(doc, "本次修订的总体原则", 2)
    for line in [
        "突出已完成、可演示、可测试、可追溯的 V0.1 产品事实。",
        "用真实测试日志、案例记录、源码仓库、演示视频和结果包支撑产品品质。",
        "把创新性落到多智能体编排、医疗科研 Skills、MCP 工具和人工门禁，而不是泛泛描述 AI。",
        "成效表述区分已验证成效和预测推广成效，避免把尚未试点的数据写成事实。",
        "安全合规使用一票否决口径，明确数据来源、脱敏、密钥和临床结论边界。",
    ]:
        add_para(doc, f"- {line}")
    doc.save(REASON_DOCX)

    lines = [
        "# 附件2申报表修正原因说明",
        "",
        "依据《副本初赛申报材料准备指南》，初赛评分重点为产品品质 25 分、创新性 30 分、应用价值 35 分、安全合规 10 分。应用价值与创新性合计 65 分，是本次修订重点。",
        "",
        "| 修正项 | 修正原因 |",
        "| --- | --- |",
    ]
    for item, reason in REASONS:
        lines.append(f"| {item} | {reason} |")
    lines.extend([
        "",
        "## 总体原则",
        "",
        "- 突出已完成、可演示、可测试、可追溯的 V0.1 产品事实。",
        "- 用真实测试日志、案例记录、源码仓库、演示视频和结果包支撑产品品质。",
        "- 把创新性落到多智能体编排、医疗科研 Skills、MCP 工具和人工门禁，而不是泛泛描述 AI。",
        "- 成效表述区分已验证成效和预测推广成效，避免把尚未试点的数据写成事实。",
        "- 安全合规使用一票否决口径，明确数据来源、脱敏、密钥和临床结论边界。",
    ])
    REASON_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    rewrite_application()
    build_reason_doc()
    print(OUTPUT)
    print(REASON_DOCX)
    print(REASON_MD)


if __name__ == "__main__":
    main()
