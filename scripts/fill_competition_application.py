"""Fill the competition application template while preserving its table layout."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path("/Users/shanbaotao/附件2：“智行合e，全员AI秀”AI应用技能大赛申报表.docx")
OUTPUT = Path(
    "/Users/shanbaotao/Documents/agent 2/deliverables/"
    "附件2-智行合e全员AI秀AI应用技能大赛申报表-临研智策-V0.1.docx"
)


FIELDS = {
    "案例名称": "临研智策：多智能体临床科研工作台",
    "个人/团队信息": "队长：山宝涛（部门、工号、岗位请按实际信息补齐）",
    "参赛方向": "技术研发与智能开发",
    "应用场景简述": (
        "面向临床科研人员的科研辅助场景。围绕课题构思、研究设计、文献检索与筛选、"
        "证据抽取与系统评价、科研写作草稿四类高频任务，形成“输入科研问题—真实工具执行—"
        "项目留痕—人工复核—受控导出”的可追溯工作流。"
    ),
    "使用AI工具/智能体类型": (
        "OpenCode 多智能体编排；基于 medical-research-skills 的 30 个筛选医疗科研 Skills；"
        "FastAPI + MCP 工具服务；React B/S 成果工作台；PubMed、Europe PMC 等公开医学数据源。"
    ),
    "改造前现状/痛点": (
        "临床科研从选题到成稿通常依赖个人经验和分散工具：研究问题、检索式、题录、证据表和"
        "写作稿件分散保存；检索与筛选过程难复现；关键方法学假设、证据来源和版本缺少统一留痕；"
        "高风险步骤容易把 AI 建议误当成最终结论。"
    ),
    "落地方案": (
        "采用“Agent 编排 + Skill 方法规范 + MCP 真实执行 + 项目事实记录 + 人工确认”的技术路线。"
        "已封装四个业务 Agent：\n"
        "1. 研究设计 Agent：生成 PICO、纳排标准、结局、基础样本量和随机化计划；\n"
        "2. 文献检索综述 Agent：生成检索式，调用 PubMed/Europe PMC，完成导入、去重、筛选建议和 PRISMA；\n"
        "3. 证据抽取系统评价 Agent：保存全文来源、抽取基线/结局、初评偏倚风险并完成限定范围 Meta；\n"
        "4. 科研写作 Agent：仅基于已保存事实输出带来源清单和待解决项的版本化草稿。\n"
        "B/S 工作台统一查看项目、审计、审批状态与受控导出结果。"
    ),
    "落地成效或预测成效": (
        "已完成 MVP 落地与功能验证，不是预测原型：后端自动化测试 39 项通过；已验证真实公开医学"
        "数据库检索、题录去重、PRISMA、公开全文入库、结构化证据抽取、RoB 2 初评、二分类 Meta、"
        "研究设计样本量计算、受控随机化及科研写作草稿。\n"
        "通过统一项目记录、Skill 执行回执、工具调用审计和人工门禁，将原本分散的科研辅助步骤"
        "收敛为可继续、可追溯、可导出的闭环，减少重复整理与交接成本。"
    ),
    "落地使用周期": "2026 年 7 月完成 V0.1 开发、真实案例验证与参赛演示材料整理；当前为本地 MVP 验证阶段。",
    "可行性评估": (
        "技术可行：复用开源 medical-research-skills、OpenCode 与 MCP 标准协议，避免重复建设；"
        "核心能力已通过自动化测试和真实公开数据案例验证。\n"
        "推广可行：采用模块化 Agent/Skill/Tool 组合，更换专科、疾病或研究问题时可复用流程与产物结构。\n"
        "资源需求：MVP 可本地运行；生产化需接入院内统一认证、权限审计、对象存储、PostgreSQL 与合规数据治理体系。"
    ),
    "数据安全与合规说明": (
        "仅处理脱敏、聚合或公开数据，禁止将患者身份信息写入模型提示词、MCP 调用或普通日志；"
        "不提供诊断、治疗建议或最终临床决策。\n"
        "项目、题录、证据、草稿、审批和导出均保留审计记录；关键步骤校验签名 Skill 执行回执。"
        "样本量假设、筛选定稿、随机化、系统评价和写作导出均保留人工复核/审批；RCT 分配序列"
        "与审批密钥不进入模型、普通网页或普通导出包。"
    ),
    "附件清单": (
        "佐证材料1-四智能体演示视频索引及 4 段操作视频；\n"
        "佐证材料2-产品截图；\n"
        "佐证材料3-产品架构与业务闭环；\n"
        "佐证材料4-真实案例与成果说明；\n"
        "佐证材料5-测试验证与数据安全合规说明；\n"
        "佐证材料6-可复用能力清单；\n"
        "佐证材料7-临床科研智能体工作台参赛汇报 PPT。"
    ),
}


def set_cell_text(cell, text):
    """Replace a template instruction cell without changing table geometry."""
    paragraph = cell.paragraphs[0]
    # The template uses several paragraphs for its grey guidance text. Remove
    # every extra paragraph so no instructions are carried into the submission.
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "SimSun"
        run.font.size = Pt(9)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE)
    table = document.tables[0]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label in FIELDS:
            set_cell_text(row.cells[1], FIELDS[label])

    document.core_properties.title = "临研智策：多智能体临床科研工作台 - AI 应用技能大赛申报表"
    document.core_properties.subject = "智行合e・全员AI秀 AI 应用技能大赛申报材料"
    document.core_properties.author = "山宝涛"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
